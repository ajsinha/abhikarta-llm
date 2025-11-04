"""
Copyright 2025-2030 all rights reserved
Ashutosh Sinha
Email: ajsinha@gmail.com
Version: 3.1.4

Document Parser - Extract text from various document formats
"""

import os
from typing import Union, Optional, Dict, Any
from pathlib import Path
from dataclasses import dataclass
from enum import Enum


class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    CSV = "csv"
    XLSX = "xlsx"
    PPTX = "pptx"
    JSON = "json"
    XML = "xml"


@dataclass
class ParsedDocument:
    """
    Represents a parsed document.
    
    Attributes:
        text: Extracted text content
        format: Document format
        metadata: Additional metadata
        page_count: Number of pages (if applicable)
        filename: Original filename
    """
    text: str
    format: DocumentFormat
    metadata: Dict[str, Any]
    page_count: Optional[int] = None
    filename: Optional[str] = None
    
    def get_summary(self, max_chars: int = 500) -> str:
        """Get summary of document"""
        if len(self.text) <= max_chars:
            return self.text
        return self.text[:max_chars] + "..."


class DocumentParser:
    """
    Parser for extracting text from various document formats.
    
    Example:
        parser = DocumentParser()
        doc = parser.parse('document.pdf')
        print(doc.text)
        print(f"Pages: {doc.page_count}")
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': DocumentFormat.PDF,
        '.docx': DocumentFormat.DOCX,
        '.doc': DocumentFormat.DOC,
        '.txt': DocumentFormat.TXT,
        '.md': DocumentFormat.MD,
        '.html': DocumentFormat.HTML,
        '.htm': DocumentFormat.HTML,
        '.csv': DocumentFormat.CSV,
        '.xlsx': DocumentFormat.XLSX,
        '.pptx': DocumentFormat.PPTX,
        '.json': DocumentFormat.JSON,
        '.xml': DocumentFormat.XML,
    }
    
    def parse(self, filepath: Union[str, Path]) -> ParsedDocument:
        """
        Parse a document and extract text.
        
        Args:
            filepath: Path to document
        
        Returns:
            ParsedDocument object
        
        Example:
            doc = parser.parse('report.pdf')
            print(doc.text)
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"Document not found: {filepath}")
        
        ext = filepath.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}")
        
        doc_format = self.SUPPORTED_FORMATS[ext]
        
        # Parse based on format
        if doc_format == DocumentFormat.PDF:
            return self._parse_pdf(filepath)
        elif doc_format == DocumentFormat.DOCX:
            return self._parse_docx(filepath)
        elif doc_format == DocumentFormat.TXT:
            return self._parse_txt(filepath)
        elif doc_format == DocumentFormat.MD:
            return self._parse_txt(filepath)  # Same as TXT
        elif doc_format == DocumentFormat.HTML:
            return self._parse_html(filepath)
        elif doc_format == DocumentFormat.CSV:
            return self._parse_csv(filepath)
        elif doc_format == DocumentFormat.XLSX:
            return self._parse_xlsx(filepath)
        elif doc_format == DocumentFormat.JSON:
            return self._parse_json(filepath)
        else:
            return self._parse_generic(filepath, doc_format)
    
    def _parse_pdf(self, filepath: Path) -> ParsedDocument:
        """Parse PDF document"""
        try:
            # Try pypdf2 first
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return ParsedDocument(
                    text=text.strip(),
                    format=DocumentFormat.PDF,
                    metadata={'tool': 'PyPDF2'},
                    page_count=len(reader.pages),
                    filename=filepath.name
                )
        except ImportError:
            # Fallback: return placeholder
            return ParsedDocument(
                text=f"[PDF Document: {filepath.name}]\nNote: Install PyPDF2 to extract text: pip install PyPDF2",
                format=DocumentFormat.PDF,
                metadata={'tool': 'placeholder'},
                filename=filepath.name
            )
    
    def _parse_docx(self, filepath: Path) -> ParsedDocument:
        """Parse DOCX document"""
        try:
            import docx
            doc = docx.Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return ParsedDocument(
                text=text.strip(),
                format=DocumentFormat.DOCX,
                metadata={'tool': 'python-docx'},
                page_count=None,
                filename=filepath.name
            )
        except ImportError:
            return ParsedDocument(
                text=f"[DOCX Document: {filepath.name}]\nNote: Install python-docx: pip install python-docx",
                format=DocumentFormat.DOCX,
                metadata={'tool': 'placeholder'},
                filename=filepath.name
            )
    
    def _parse_txt(self, filepath: Path) -> ParsedDocument:
        """Parse text file"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return ParsedDocument(
            text=text,
            format=DocumentFormat.TXT,
            metadata={'encoding': 'utf-8'},
            filename=filepath.name
        )
    
    def _parse_html(self, filepath: Path) -> ParsedDocument:
        """Parse HTML file"""
        try:
            from bs4 import BeautifulSoup
            with open(filepath, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                text = soup.get_text()
            
            return ParsedDocument(
                text=text.strip(),
                format=DocumentFormat.HTML,
                metadata={'tool': 'beautifulsoup4'},
                filename=filepath.name
            )
        except ImportError:
            # Fallback: basic HTML parsing
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            return ParsedDocument(
                text=text,
                format=DocumentFormat.HTML,
                metadata={'tool': 'basic'},
                filename=filepath.name
            )
    
    def _parse_csv(self, filepath: Path) -> ParsedDocument:
        """Parse CSV file"""
        import csv
        rows = []
        with open(filepath, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))
        
        text = "\n".join(rows)
        
        return ParsedDocument(
            text=text,
            format=DocumentFormat.CSV,
            metadata={'rows': len(rows)},
            filename=filepath.name
        )
    
    def _parse_xlsx(self, filepath: Path) -> ParsedDocument:
        """Parse Excel file"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = ", ".join([str(cell) if cell is not None else "" for cell in row])
                    text_parts.append(row_text)
                text_parts.append("")
            
            return ParsedDocument(
                text="\n".join(text_parts).strip(),
                format=DocumentFormat.XLSX,
                metadata={'sheets': len(wb.sheetnames)},
                filename=filepath.name
            )
        except ImportError:
            return ParsedDocument(
                text=f"[XLSX Document: {filepath.name}]\nNote: Install openpyxl: pip install openpyxl",
                format=DocumentFormat.XLSX,
                metadata={'tool': 'placeholder'},
                filename=filepath.name
            )
    
    def _parse_json(self, filepath: Path) -> ParsedDocument:
        """Parse JSON file"""
        import json
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        text = json.dumps(data, indent=2)
        
        return ParsedDocument(
            text=text,
            format=DocumentFormat.JSON,
            metadata={'type': 'json'},
            filename=filepath.name
        )
    
    def _parse_generic(self, filepath: Path, doc_format: DocumentFormat) -> ParsedDocument:
        """Generic parser for unsupported formats"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except:
            text = f"[Binary file: {filepath.name}]"
        
        return ParsedDocument(
            text=text,
            format=doc_format,
            metadata={'parser': 'generic'},
            filename=filepath.name
        )


__all__ = [
    'DocumentFormat',
    'ParsedDocument',
    'DocumentParser',
]


"""
Copyright 2025-2030 all rights reserved
Ashutosh Sinha | ajsinha@gmail.com | Version: 3.1.4
"""
