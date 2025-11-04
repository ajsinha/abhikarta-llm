"""
Copyright 2025-2030 all rights reserved
Ashutosh Sinha
Email: ajsinha@gmail.com
Version: 3.1.4

File Processing - Transparent handling of PDF, DOCX, CSV, TXT, JSON
"""

import json
import csv
import io
from typing import Optional, Dict, Any, List, Union
from pathlib import Path
from dataclasses import dataclass


@dataclass
class ProcessedFile:
    """
    Processed file result.
    
    Attributes:
        content: Extracted text content
        metadata: File metadata
        file_type: Type of file
        structured_data: For CSV/JSON - structured data
    """
    content: str
    metadata: Dict[str, Any]
    file_type: str
    structured_data: Optional[Union[List, Dict]] = None


class PDFProcessor:
    """Process PDF files"""
    
    @staticmethod
    def extract_text(filepath: str) -> ProcessedFile:
        """
        Extract text from PDF.
        
        Args:
            filepath: Path to PDF file
        
        Returns:
            ProcessedFile with extracted text
        """
        try:
            # Try using pypdf2
            import PyPDF2
            
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                metadata = {
                    "pages": len(reader.pages),
                    "filename": Path(filepath).name
                }
                
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return ProcessedFile(
                    content=text.strip(),
                    metadata=metadata,
                    file_type="pdf"
                )
        except ImportError:
            # Fallback: basic text extraction
            return ProcessedFile(
                content=f"PDF file: {filepath}\n(Install PyPDF2 for text extraction)",
                metadata={"filename": Path(filepath).name},
                file_type="pdf"
            )
        except Exception as e:
            return ProcessedFile(
                content=f"Error processing PDF: {e}",
                metadata={"filename": Path(filepath).name, "error": str(e)},
                file_type="pdf"
            )
    
    @staticmethod
    def extract_with_layout(filepath: str) -> ProcessedFile:
        """Extract PDF preserving layout"""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n\n"
            
            return ProcessedFile(
                content=text.strip(),
                metadata={"filename": Path(filepath).name},
                file_type="pdf"
            )
        except ImportError:
            return PDFProcessor.extract_text(filepath)


class DOCXProcessor:
    """Process DOCX files"""
    
    @staticmethod
    def extract_text(filepath: str) -> ProcessedFile:
        """
        Extract text from DOCX.
        
        Args:
            filepath: Path to DOCX file
        
        Returns:
            ProcessedFile with extracted text
        """
        try:
            from docx import Document
            
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            metadata = {
                "filename": Path(filepath).name,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return ProcessedFile(
                content=text.strip(),
                metadata=metadata,
                file_type="docx"
            )
        except ImportError:
            return ProcessedFile(
                content=f"DOCX file: {filepath}\n(Install python-docx for text extraction)",
                metadata={"filename": Path(filepath).name},
                file_type="docx"
            )
        except Exception as e:
            return ProcessedFile(
                content=f"Error processing DOCX: {e}",
                metadata={"filename": Path(filepath).name, "error": str(e)},
                file_type="docx"
            )
    
    @staticmethod
    def extract_with_tables(filepath: str) -> ProcessedFile:
        """Extract DOCX including tables"""
        try:
            from docx import Document
            
            doc = Document(filepath)
            content_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            for table in doc.tables:
                table_text = "\n[TABLE]\n"
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    table_text += " | ".join(cells) + "\n"
                content_parts.append(table_text)
            
            return ProcessedFile(
                content="\n".join(content_parts),
                metadata={"filename": Path(filepath).name},
                file_type="docx"
            )
        except:
            return DOCXProcessor.extract_text(filepath)


class CSVProcessor:
    """Process CSV files"""
    
    @staticmethod
    def read_csv(filepath: str, max_rows: Optional[int] = None) -> ProcessedFile:
        """
        Read CSV file.
        
        Args:
            filepath: Path to CSV file
            max_rows: Maximum rows to read
        
        Returns:
            ProcessedFile with CSV data
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                
                if max_rows:
                    rows = rows[:max_rows]
                
                # Create text representation
                text_parts = []
                text_parts.append(f"CSV File: {Path(filepath).name}")
                text_parts.append(f"Columns: {', '.join(rows[0].keys() if rows else [])}")
                text_parts.append(f"Rows: {len(rows)}")
                text_parts.append("\nData:")
                
                for i, row in enumerate(rows[:10]):  # First 10 rows
                    text_parts.append(f"Row {i+1}: {row}")
                
                if len(rows) > 10:
                    text_parts.append(f"... and {len(rows) - 10} more rows")
                
                return ProcessedFile(
                    content="\n".join(text_parts),
                    metadata={
                        "filename": Path(filepath).name,
                        "rows": len(rows),
                        "columns": list(rows[0].keys() if rows else [])
                    },
                    file_type="csv",
                    structured_data=rows
                )
        except Exception as e:
            return ProcessedFile(
                content=f"Error processing CSV: {e}",
                metadata={"filename": Path(filepath).name, "error": str(e)},
                file_type="csv"
            )
    
    @staticmethod
    def to_markdown_table(filepath: str, max_rows: int = 100) -> str:
        """Convert CSV to markdown table"""
        result = CSVProcessor.read_csv(filepath, max_rows)
        if not result.structured_data:
            return result.content
        
        rows = result.structured_data
        if not rows:
            return "Empty CSV"
        
        # Create markdown table
        headers = list(rows[0].keys())
        lines = []
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---" for _ in headers]) + "|")
        
        for row in rows:
            lines.append("| " + " | ".join(str(row.get(h, "")) for h in headers) + " |")
        
        return "\n".join(lines)


class TXTProcessor:
    """Process TXT files"""
    
    @staticmethod
    def read_text(filepath: str, encoding: str = 'utf-8') -> ProcessedFile:
        """
        Read text file.
        
        Args:
            filepath: Path to text file
            encoding: File encoding
        
        Returns:
            ProcessedFile with text content
        """
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            
            lines = content.split('\n')
            words = content.split()
            
            return ProcessedFile(
                content=content,
                metadata={
                    "filename": Path(filepath).name,
                    "lines": len(lines),
                    "words": len(words),
                    "characters": len(content)
                },
                file_type="txt"
            )
        except Exception as e:
            return ProcessedFile(
                content=f"Error processing TXT: {e}",
                metadata={"filename": Path(filepath).name, "error": str(e)},
                file_type="txt"
            )


class JSONProcessor:
    """Process JSON files"""
    
    @staticmethod
    def read_json(filepath: str) -> ProcessedFile:
        """
        Read JSON file.
        
        Args:
            filepath: Path to JSON file
        
        Returns:
            ProcessedFile with JSON data
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Create text representation
            text_parts = []
            text_parts.append(f"JSON File: {Path(filepath).name}")
            text_parts.append(f"Type: {type(data).__name__}")
            
            if isinstance(data, dict):
                text_parts.append(f"Keys: {', '.join(data.keys())}")
            elif isinstance(data, list):
                text_parts.append(f"Items: {len(data)}")
            
            text_parts.append("\nContent:")
            text_parts.append(json.dumps(data, indent=2, ensure_ascii=False))
            
            return ProcessedFile(
                content="\n".join(text_parts),
                metadata={
                    "filename": Path(filepath).name,
                    "type": type(data).__name__
                },
                file_type="json",
                structured_data=data
            )
        except Exception as e:
            return ProcessedFile(
                content=f"Error processing JSON: {e}",
                metadata={"filename": Path(filepath).name, "error": str(e)},
                file_type="json"
            )


class UniversalFileProcessor:
    """
    Universal file processor - automatically detects and processes files.
    
    Example:
        processor = UniversalFileProcessor()
        result = processor.process("document.pdf")
        print(result.content)
    """
    
    def __init__(self):
        self.processors = {
            '.pdf': PDFProcessor.extract_text,
            '.docx': DOCXProcessor.extract_text,
            '.doc': DOCXProcessor.extract_text,
            '.csv': CSVProcessor.read_csv,
            '.txt': TXTProcessor.read_text,
            '.json': JSONProcessor.read_json,
            '.md': TXTProcessor.read_text,
            '.log': TXTProcessor.read_text,
        }
    
    def process(self, filepath: str) -> ProcessedFile:
        """
        Process file based on extension.
        
        Args:
            filepath: Path to file
        
        Returns:
            ProcessedFile with extracted content
        """
        path = Path(filepath)
        ext = path.suffix.lower()
        
        if ext in self.processors:
            return self.processors[ext](filepath)
        else:
            # Try as text
            try:
                return TXTProcessor.read_text(filepath)
            except:
                return ProcessedFile(
                    content=f"Unsupported file type: {ext}",
                    metadata={"filename": path.name, "extension": ext},
                    file_type="unknown"
                )
    
    def process_for_llm(self, filepath: str, max_length: int = 50000) -> str:
        """
        Process file and prepare for LLM input.
        
        Args:
            filepath: Path to file
            max_length: Maximum content length
        
        Returns:
            Formatted text for LLM
        """
        result = self.process(filepath)
        
        content = f"File: {result.metadata.get('filename', filepath)}\n"
        content += f"Type: {result.file_type}\n"
        content += f"\nContent:\n{result.content}"
        
        if len(content) > max_length:
            content = content[:max_length] + f"\n\n... (truncated, total: {len(result.content)} chars)"
        
        return content
    
    def batch_process(self, filepaths: List[str]) -> List[ProcessedFile]:
        """Process multiple files"""
        return [self.process(fp) for fp in filepaths]


def process_file(filepath: str) -> ProcessedFile:
    """
    Convenience function to process any file.
    
    Args:
        filepath: Path to file
    
    Returns:
        ProcessedFile with content
    
    Example:
        result = process_file("document.pdf")
        print(result.content)
    """
    processor = UniversalFileProcessor()
    return processor.process(filepath)


def process_file_for_prompt(filepath: str, prompt: str) -> str:
    """
    Process file and combine with prompt.
    
    Args:
        filepath: Path to file
        prompt: User prompt
    
    Returns:
        Combined prompt with file content
    
    Example:
        combined = process_file_for_prompt("data.csv", "Summarize this data")
    """
    processor = UniversalFileProcessor()
    file_content = processor.process_for_llm(filepath)
    
    return f"{prompt}\n\n{file_content}"


__all__ = [
    'ProcessedFile',
    'PDFProcessor',
    'DOCXProcessor',
    'CSVProcessor',
    'TXTProcessor',
    'JSONProcessor',
    'UniversalFileProcessor',
    'process_file',
    'process_file_for_prompt',
]


"""
Copyright 2025-2030 all rights reserved
Ashutosh Sinha | ajsinha@gmail.com | Version: 3.1.4
"""
